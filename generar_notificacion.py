import os
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def limpiar_articulo_excel(valor):
    """Convierte valores como 87.0 en '87' y maneja valores vacíos."""
    if valor is None or str(valor).lower() == 'nan':
        return ""
    val_str = str(valor)
    if val_str.endswith('.0'):
        return val_str[:-2]
    return val_str

def limpiar_fecha_excel(valor):
    """Limpia y formatea fechas desde Excel."""
    if pd.isna(valor) or str(valor).lower() == 'nan' or valor == "":
        return ""
    try:
        if isinstance(valor, (datetime, pd.Timestamp)):
            return valor.strftime("%d/%m/%Y")
        fecha_dt = pd.to_datetime(valor, dayfirst=True)
        return fecha_dt.strftime("%d/%m/%Y")
    except Exception:
        return str(valor)

def generar_notificacion_baja_word(datos, plantilla_path="plantilla_bajas.docx"):
    """Rellena la plantilla de Word con los datos del diccionario 'datos'."""
    
    if not os.path.exists(plantilla_path):
        print(f"⚠ Error: No se encuentra la plantilla {plantilla_path}")
        return

    try:
        doc = Document(plantilla_path)
        
        raw_trayecto = str(datos.get('AÑO', '')).strip().upper()
        if raw_trayecto == 'NAN' or not raw_trayecto:
            texto_trayecto = ""
        else:
            texto_trayecto = f"de {raw_trayecto} "
        
        # Diccionario de reemplazos
        reemplazos = {
            "{{NOMBRE}}": str(datos.get('NOMBRES', '')).upper(),
            "{{APELLIDO}}": str(datos.get('APELLIDO 1', '')).upper(),
            "{{CEDULA}}": str(datos.get('CÉDULA', '')),
            
            # Específicos de PNF
            "{{EJE}}": str(datos.get('EJE', '')).upper(),
            "{{ASIC}}": str(datos.get('ASIC', '')).upper(),
            
            # Específicos de PNFA
            "{{HOSPITAL}}": str(datos.get('HOSPITAL SEDE', '')).upper(),
            
            # Variables compartidas (Busca en PNF y si no, en PNFA)
            "{{TRAYECTO}}": texto_trayecto, 
            "{{CAUSAL}}": str(datos.get('CAUSAL', datos.get('MOTIVO', ''))).upper(),
            "{{FECHA_TRAMITE}}": limpiar_fecha_excel(datos.get('FECHA TRAMITE', datos.get('FECHA SOLICITUD'))),
            
            # Soportar ambas etiquetas de programa
            "{{PNF}}": str(datos.get('PNF', datos.get('PNFA', ''))).upper(),
            "{{PNFA}}": str(datos.get('PNF', datos.get('PNFA', ''))).upper(),
            
            "{{CABES}}": str(datos.get('CABES', '')).upper(),
            "{{ARTICULO}}": limpiar_articulo_excel(datos.get('ARTICULO')),
            "{{FECHA_CABES}}": limpiar_fecha_excel(datos.get('FECHA', datos.get('FECHA CABES'))),
        }
        
        def reemplazar_texto_preservando_formato(parrafo):
            """Reemplaza texto iterando sobre los runs para intentar mantener negritas."""
            full_text = parrafo.text
            match_found = False
            for key in reemplazos.keys():
                if key in full_text:
                    match_found = True
                    break
            
            if match_found:
                for key, value in reemplazos.items():
                    if key in parrafo.text:
                        parrafo.text = parrafo.text.replace(key, value)
                
                # Reaplicar fuente Calibri 10
                for run in parrafo.runs:
                    run.font.name = 'Calibri'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                    run.font.size = Pt(10)

        # 1. Reemplazar en párrafos
        for p in doc.paragraphs:
            reemplazar_texto_preservando_formato(p)

        # 2. Reemplazar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        reemplazar_texto_preservando_formato(p)

        # --- SECCIÓN DE NOMBRE DE ARCHIVO ---
        fecha_hoy = datetime.now().strftime("%d-%m-%Y")
        
        # 1. Obtener Nombre y Apellido limpios
        raw_nombre = str(datos.get('NOMBRES', 'Estudiante')).strip().upper()
        raw_apellido = str(datos.get('APELLIDO 1', '')).strip().upper()
        if not raw_apellido:
            raw_apellido = str(datos.get('APELLIDOS', '')).strip().upper()
            
        # 2. Obtener Cédula
        raw_cedula = str(datos.get('CÉDULA', '')).strip()
        if not raw_cedula or raw_cedula.lower() == 'nan':
            raw_cedula = str(datos.get('cedula', 'SN')).strip()

        # 3. Limpieza de caracteres prohibidos en nombres de archivo
        caracteres_prohibidos = ['/', '\\', ':', '*', '?', '"', '<', '>', '|']
        for char in caracteres_prohibidos:
            raw_nombre = raw_nombre.replace(char, '')
            raw_apellido = raw_apellido.replace(char, '')
            raw_cedula = raw_cedula.replace(char, '')

        # 4. Construir el nombre final INCLUYENDO la fecha de hoy
        nombre_salida = f"Notificacion_{raw_nombre}_{raw_apellido}_{raw_cedula}_{fecha_hoy}.docx"
        # --- FIN SECCIÓN MODIFICADA ---

        if not os.path.exists("Notificaciones"):
            os.makedirs("Notificaciones")
            
        ruta_salida = os.path.join("Notificaciones", nombre_salida)
        doc.save(ruta_salida)
        print(f"   Word generado: {ruta_salida}")

    except Exception as e:
        print(f"   ⚠ Error generando Word: {e}")
